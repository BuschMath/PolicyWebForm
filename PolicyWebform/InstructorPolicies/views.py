from django.shortcuts import render
from .forms import DocumentForm
from docx import Document

# Create your views here.
def create_document(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST)
        if form.is_valid():
            doc = Document()
            doc.add_paragraph('Instructor Name:', form.cleaned_data['InstructorName'])
            doc.add_paragraph('Course Name and Section:', form.cleaned_data['CourseNameSection'])
            doc.add_paragraph('Office Location:', form.cleaned_data['OfficeLocation'])
            doc.add_paragraph('Phone:', form.cleaned_data['Phone'])
            doc.add_paragraph('Email:', form.cleaned_data['Email'])
            doc.add_paragraph('Office Hours:', form.cleaned_data['OfficeHours'])
            doc.add_paragraph('Final Exam Date:', form.cleaned_data['FinalExamDate'])
            doc.add_paragraph('Final Exam Time:', form.cleaned_data['FinalExamTime'])
            doc.add_paragraph('Textbook:', form.cleaned_data['Textbook'])
            doc.add_paragraph('Appropriate Behavior:', form.cleaned_data['AppropriateBehavior'])
            doc.add_paragraph('Assessments:', form.cleaned_data['Assessments'])
            doc.add_paragraph('Grading Scale:', form.cleaned_data['GradingScale'])
            doc.add_paragraph('Late Work:', form.cleaned_data['LateWork'])
            doc.add_paragraph('Extra Credit:', form.cleaned_data['ExtraCredit'])
            doc.add_paragraph('Final Exam Policy:', form.cleaned_data['FinalExamPolicy'])
            doc.add_paragraph('Required Materials:', form.cleaned_data['RequiredMaterials'])
            doc.add_paragraph('Suggested Materials:', form.cleaned_data['SuggestedMaterials'])
            doc.add_paragraph('Course Fees:', form.cleaned_data['CourseFees'])
            doc.add_paragraph('Group Work:', form.cleaned_data['GroupWork'])
            doc.add_paragraph('Previous Work:', form.cleaned_data['PreviousWork'])
            doc.add_paragraph('Scholastic Honesty:', form.cleaned_data['ScholasticHonesty'])
            doc.add_paragraph('Attendance Policy:', form.cleaned_data['AttendancePolicy'])
            doc.add_paragraph('Communication Policy:', form.cleaned_data['CommunicationPolicy'])
            doc.add_paragraph('Appropriate Technology Use:', form.cleaned_data['AppropriateTechnologyUse'])
            doc.add_paragraph('Personal Responsibility:', form.cleaned_data['PersonalResponsibility'])
            doc.add_paragraph('Standards For Written Work:', form.cleaned_data['StandardsForWrittenWork'])
            doc.add_paragraph('Computer Considerations:', form.cleaned_data['ComputerConsiderations'])
            doc.add_paragraph('Online Learning Platform:', form.cleaned_data['OnlineLearningPlatform'])
            doc.add_paragraph('Teaching Philosophy:', form.cleaned_data['TeachingPhilosophy'])

            doc.save('output.docx')

    else:
        form = DocumentForm()

    return render(request, 'InstructorPolicies/create_document.html', {'form': form})
